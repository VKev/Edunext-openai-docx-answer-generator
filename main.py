from docx import Document
from docx.table import Table
from docx.shared import Pt

from openai import OpenAI

import time
from selenium import webdriver
from selenium.webdriver.chrome.service import Service 
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait as wait
import secrets
import copy

import random
import os

import requests
import re
from transformers import PegasusForConditionalGeneration, PegasusTokenizer
from transformers import logging
logging.set_verbosity_error()

def paraphrase_text(text, model_name="tuner007/pegasus_paraphrase"):
    # Load the tokenizer and model
    tokenizer = PegasusTokenizer.from_pretrained(model_name)
    model = PegasusForConditionalGeneration.from_pretrained(model_name)

    # Split the text into sentences based on special symbols
    sentences = re.split(r'(?<=[.!?]) +', text)

    # Paraphrase each sentence
    paraphrased_sentences = []
    for sentence in sentences:
        # Skip any empty sentences
        if sentence.strip():
            tokens = tokenizer(sentence, truncation=True, padding="longest", return_tensors="pt")
            outputs = model.generate(**tokens, num_beams=5, num_return_sequences=1, max_length=100)
            paraphrased_sentence = tokenizer.decode(outputs[0], skip_special_tokens=True)
            paraphrased_sentences.append(paraphrased_sentence)

    # Join the paraphrased sentences
    paraphrased_text = ' '.join(paraphrased_sentences)
    return paraphrased_text


def get_google_link_by_string(search_string):
    url = 'https://www.googleapis.com/customsearch/v1'
    params = {
        'q': search_string,
        'key': os.environ.get("GOOGLE_API_KEY"),
        'cx': '216b6900a99d943fb',
        'num': 10  # Adjust 'num' to fetch more results, up to a maximum of 10 per page.
    }
    response = requests.get(url, params=params)
    result = response.json()
    if 'items' in result:

        random_index = random.randint(0, 2)
        return result['items'][random_index]['link']
    else:
        return None

def getFilePath(filename):
    return os.path.join(current_directory, filename)

def append_content(from_doc, to_doc):
    for element in from_doc.element.body:
        to_doc.element.body.append(copy.deepcopy(element))
def get_number_of_lines(doc):
    return len(doc.element.body)
def append_para_by_lineIndex(document,content, lineIndex):
    paragraph = document.paragraphs[lineIndex]
    new_run = paragraph.add_run()
    old_run = paragraph.runs[0]
    new_run.text += content
    new_run.bold = old_run.bold
    new_run.italic = old_run.italic
    new_run.underline = old_run.underline
    new_run.font.color.rgb = old_run.font.color.rgb
    new_run.font.name = old_run.font.name
    new_run.font.size = old_run.font.size
def append_to_row(table, row_index, string):
    if row_index < len(table.rows):
        cell = table.cell(row_index, 0) 
        cell.text += string
    else:
        print("Row index out of range.")


def append_to_row_custom(table, row_index, gpt_answer, reference_link):
    if row_index < len(table.rows):
        cell = table.cell(row_index, 0)
        p = cell.paragraphs[0]
        p.add_run(gpt_answer + '\n\n')
        run = p.add_run('Reference: ')
        run.bold = True
        run = p.add_run(reference_link)
        run.hyperlink = reference_link
    else:
        print("Row index out of range.")

current_directory = os.path.dirname(os.path.abspath(__file__))

client = OpenAI(
    api_key=os.environ.get("OPENAI_API_KEY"),
)


tittleTemplate = Document(getFilePath('edunextTemplate_Tittle.docx'))
questionTemplate = Document(getFilePath('edunextTemplate_Question.docx'))
conclusionTemplate = Document(getFilePath('edunextTemplate_Conclusion.docx'))



def create_link(answer: str):
    link = "https://www.example.com/search?q=" + answer.replace(" ", "+")
    return link

keyword = input("Enter any relate keyword: ")
def create_gpt_response(question: str, maxtoken : int):
    unique_seed = random.randint(1, 1000000)  # Generate a random seed
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"{question} in {keyword}. (unique seed: {unique_seed}).",
            }
        ],
        model="gpt-3.5-turbo-0125",
        max_tokens=  maxtoken,
    )
    return chat_completion.choices[0].message.content


option = webdriver.ChromeOptions()

option.add_argument("user-data-dir="+'C:\\MyLaptop\\VisualStudioCode\\GitProject\\EdunextAutomatic\\chromeDriver\\localhost')

class Brower:
    driver, service = None, None
    def __init__(self, chromeDriver: str):
        self.service = Service(chromeDriver)
        self.driver = webdriver.Chrome(service=self.service, options= option)

    def open_page(self,url: str):
        self.driver.get(url)

    def close_page(self):
        self.driver.close()

    def click_button_by_xpath(self, xpath:str):
        self.driver.find_element(By.XPATH, xpath).click()
        time.sleep(1.5)
    
    def send_data_to_xpath(self,xpath:str, data:str):
        self.driver.find_element(By.XPATH, xpath).send_keys(data)
        time.sleep(1.5)

waittimeEachGPTresponse = input("Enter wait time each gpt response: ")

if __name__ == '__main__':
    browser = Brower('chromeDriver\chromedriver.exe')
    browser.open_page('https://fu-edunext.fpt.edu.vn/course?id=889&classId=5189')
    nextpage = input('Next page? y/n:')
    if(nextpage == 'y'):
        browser.click_button_by_xpath('//button[@class="paging-button ml-0"]')
    else:
        browser.click_button_by_xpath('//button[contains(text(),"Prev")]')


    parent_content = browser.driver.find_element(By.XPATH, "//div[@class='container-sessions']")
    all_children_by_xpath = parent_content.find_elements(By.XPATH,"./*")
    ii=1
    startSlot = int(input("Enter start slot: "))
    endSlot = int(input("Enter end slot: "))
    cooldown = int(input("Enter cool down: "))
    
    for child in all_children_by_xpath:

        if(ii<startSlot): 
            ii+=1
            continue
        if(ii > endSlot):
            break  
            
        filename = "result"+str(ii)+".docx"
        document = Document()
        append_content(tittleTemplate, document)
        append_para_by_lineIndex(document, ' '+str(ii), 7 )

        print("\nSLOT "+str(ii)+":")
        questions_a_tag = child.find_elements(By.XPATH, ".//a[contains(@class, 'mg-b-0 text-normal text-decoration-none')]")
        links = [element.get_attribute('href') for element in questions_a_tag]

        currentQuestionLineIndex = 9
        currentAnswerTableIndex = 0

        edunext_questions_tab = browser.driver.window_handles[0]
        questionList = []
        print('Progress: ', end = '')
        for i in range(len(links)):
            if(i==0):
                browser.driver.execute_script('window.open(arguments[0]);', links[i])
                browser.driver.switch_to.window(browser.driver.window_handles[1])  
                time.sleep(1.5)
            else:
                browser.driver.get(links[i])
                time.sleep(1.5)
                currentQuestionLineIndex += 7
                currentAnswerTableIndex+=5

            all_question_tag = browser.driver.find_elements(By.XPATH, "//div[@class='styled']//*")
            texts = [child.text for child in all_question_tag]
            text = '\n'.join(texts) 
            questionList.append(text)
            append_content(questionTemplate, document)
            
            for tableIndex in range(currentAnswerTableIndex, currentAnswerTableIndex+5):
                gpt_answer = create_gpt_response(text, random.randint(100, 200))
                time.sleep(int(waittimeEachGPTresponse))
                reference_link = get_google_link_by_string(text)
                if(reference_link==''):
                    reference_link = get_google_link_by_string(keyword)
                append_to_row_custom(document.tables[tableIndex], 1, gpt_answer, reference_link)
                print(str( tableIndex+1)+" ", end = '')
            append_para_by_lineIndex(document, str(i+1)+'.'+ text, currentQuestionLineIndex)
            #print("QQQQQQQQQQQQQQQ\n"+create_gpt_response(text)+'\n')
        print()
        conclusionDoc = Document()
        qindex = 1
        append_content(conclusionTemplate, document)
        for q in questionList:
            newpara = conclusionDoc.add_paragraph()
            run = newpara.add_run('Q' + str(qindex) + ':' + q)
            run.bold = True
            gpt_answer = create_gpt_response(q+ 'unique seed: ' + str(random.randint(1, 1000000)), random.randint(350, 400))
            newpara.add_run('\n'+gpt_answer)
            reference_link = get_google_link_by_string(text)
            if(reference_link == ''):
                reference_link = get_google_link_by_string(keyword)
            run = newpara.add_run("\nReference: ")
            run.bold = True
            run = newpara.add_run(reference_link)
            qindex += 1

        append_content(conclusionDoc, document)
        
        browser.driver.close()
        browser.driver.switch_to.window(edunext_questions_tab)
        
        document.save(getFilePath(filename))   
        time.sleep(cooldown)
        ii+=1
    
    browser.close_page()





